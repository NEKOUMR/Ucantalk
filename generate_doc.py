"""
Generate the Ucantalk manual docx from an existing template document.
"""
import argparse
import copy
import os
from pathlib import Path
import sys

sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

PROJECT_ROOT = Path(__file__).resolve().parent


def parse_args():
    parser = argparse.ArgumentParser(description='Generate the Ucantalk manual docx.')
    parser.add_argument(
        '--template',
        default='',
        help='Path to the source docx template. Can also be provided via UCANTALK_DOC_TEMPLATE.',
    )
    parser.add_argument(
        '--output',
        default=str(PROJECT_ROOT / 'Ucantalk使用手册.docx'),
        help='Output docx path. Defaults to Ucantalk使用手册.docx in the project root.',
    )
    return parser.parse_args()


args = parse_args()
template_arg = args.template.strip()
template_value = template_arg or os.environ.get('UCANTALK_DOC_TEMPLATE', '').strip()
if not template_value:
    raise SystemExit('Template docx path is required. Pass --template <path> or set UCANTALK_DOC_TEMPLATE.')

TEMPLATE = str(Path(template_value).expanduser().resolve())
OUTPUT = str(Path(args.output).expanduser().resolve())

GRAY = RGBColor(0x93, 0x93, 0x93)
BLACK = RGBColor(0x00, 0x00, 0x00)

doc = Document(TEMPLATE)

# Clear all body content (keep styles)
body = doc.element.body
for child in list(body):
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag in ('p', 'tbl', 'sdt'):
        body.remove(child)

# ── Helpers ──────────────────────────────────────────────────────────────────

def h1(text):
    p = doc.add_paragraph(style='Heading 1')
    p.add_run(text)
    return p

def h2(text):
    p = doc.add_paragraph(style='Heading 2')
    p.add_run(text)
    return p

def h3(text):
    p = doc.add_paragraph(style='Heading 3')
    p.add_run(text)
    return p

def h4(text):
    p = doc.add_paragraph(style='Heading 4')
    p.add_run(text)
    return p

def normal(*parts):
    """parts: list of (text, bold=False) tuples or plain strings"""
    p = doc.add_paragraph(style='Normal')
    for part in parts:
        if isinstance(part, str):
            p.add_run(part)
        else:
            text, bold = part
            r = p.add_run(text)
            r.bold = bold
    return p

def gray_small(*parts):
    """Gray small text for metadata lines"""
    p = doc.add_paragraph(style='Normal')
    for part in parts:
        if isinstance(part, str):
            r = p.add_run(part)
        else:
            text, bold = part
            r = p.add_run(text)
            r.bold = bold
        r.font.size = Pt(11)
        r.font.color.rgb = GRAY
    return p

def blank():
    doc.add_paragraph(style='Normal')

def b(text):
    return (text, True)

def n(text):
    return (text, False)

# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT CONTENT
# ═══════════════════════════════════════════════════════════════════════════════

h1('📘 Ucantalk 使用手册')
gray_small('作者: NEKO_UMR')
gray_small('日期: 2026年3月')
gray_small('版本: C# 重构版（基于 WinUI 3 / .NET 8）')
gray_small('适用系统: Windows 10 (17763) 及以上，支持 x86 / x64 / ARM64')

blank()

normal(b('📢 交流与反馈'))
normal(b('QQ群号: 1040186268'), n(' (不懂的问题或Bug反馈请加群)'))
normal(b('提示: '), n('遇到问题请优先查看本手册底部的 '), b('[常见问题汇总]'), n('，实在搞不懂再来问群主。'))
normal(n('可以把使用手册下载一个副本丢给AI，让AI解决你的问题。'))

blank()

# ─── 第一章 ────────────────────────────────────────────────────────────────────
h2('📦 第一章: 安装与初始化')
normal(n('与旧版（Python版）最大的不同是：'), b('新版已完全不需要 Python 环境，也不需要手动安装插件。'), n('所有功能已内置，开箱即用。'))

h3('1. 安装软件')
normal(n('根据下载包类型，二选一：'))
normal(b('方式A：安装包（推荐）'))
normal(n('运行 '), b('Ucantalk_Setup.exe'), n('，按提示完成安装即可。'))
normal(n('程序会自动安装 .NET 8 运行库和 VC++ 依赖。'))
normal(n('安装完成后，从开始菜单或桌面快捷方式启动。'))
normal(b('方式B：MSIX 包'))
normal(n('双击 '), b('.msix'), n(' 文件，点击"安装"。'))
normal(n('如提示证书问题，先双击 '), b('.cer'), n(' 文件安装证书，再安装 MSIX。'))
normal(b('⚠️ 建议: '), n('右键软件图标 -> 属性 -> 兼容性 -> 勾选 '), b('"以管理员身份运行此程序"'), n(' (防止在游戏全屏时全局热键失效)。'))

h3('2. 安装虚拟声卡 (仅首次，VRChat 用户必须做)')
normal(n('在软件下载包内，找到名为 '), b('VBCABLE_Driver_Pack45'), n(' (或类似名称) 的文件夹。'))
normal(n('右键点击 '), b('VBCABLE_Setup_x64.exe'), n('，选择 '), b('[以管理员身份运行]'), n('。'))
normal(n('点击 '), b('"Install Driver"'), n('。'))
normal(b('⚠️ 关键步骤: 安装完成后，必须重启电脑!'))

h3('3. 检查 FFmpeg (已内置)')
normal(n('软件目录下的 '), b('tools/ffmpeg/'), n(' 文件夹中已自带 '), b('ffmpeg.exe'), n('。'))
normal(b('注意: '), n('请勿删除或移动该文件，否则 '), b('音频播放器'), n(' 和 '), b('语音合成'), n(' 将无法工作。'))

h3('4. 语音识别模型 (已内置)')
normal(n('新版已在 '), b('tools/sherpa-default/'), n(' 内置了高质量中文离线语音识别模型，'), b('无需额外下载'), n('。'))
normal(n('首次使用语音识别时，直接在语音识别设置中选择 '), b('"Sherpa-ONNX（默认模型）"'), n(' 即可。'))

h3('5. 启动软件')
normal(n('从开始菜单或桌面快捷方式双击启动 '), b('Ucantalk'), n('。'))
normal(b('建议: '), n('右键 -> 属性 -> 兼容性 -> 勾选 '), b('"以管理员身份运行此程序"'), n(' (防止在游戏全屏时快捷键失效)。'))

blank()

# ─── 第二章 ────────────────────────────────────────────────────────────────────
h2('🎛️ 第二章: 音频设置 (关键步骤)')
normal(n('新版自带 '), b('"双通道音频路由"'), n('，能自动把声音同时发给你和 VRChat，无需由 OBS 中转。'))

h3('2.1 软件内部设置')
normal(n('打开软件，点击左侧导航栏的 '), b('⚙️ 设置'), n(' -> '), b('🔊 音频路由'), n(' 区域:'))
normal(b('🎧 监听设备 (给自己听):'))
normal(b('操作: '), n('点击下拉框，选择你的物理耳机/音响 (例如 Speakers (Realtek Audio)).'))
normal(b('效果: '), n('你能实时听到 AI 说话和播放器的音乐。'))
normal(b('📡 VRC设备 (给麦克风):'))
normal(b('操作: '), n('点击下拉框，选择 '), b('CABLE Input (VB-Audio Virtual Cable)'), n('.'))
normal(b('效果: '), n('将纯净的声音发送给游戏。'))
normal(b('❓ 找不到设备? '), n('点击右侧的 '), b('🔄 刷新'), n(' 按钮。如果还是没有 CABLE Input，说明第一章的驱动没装好或没重启。'))

h3('2.2 VRChat 游戏设置')
normal(n('进入 VRChat 菜单 -> '), b('Audio'), n(' -> '), b('Microphone'), n('。'))
normal(n('输入设备选择: '), b('CABLE Output (VB-Audio Virtual Cable)'), n('.'))

blank()

# ─── 第三章 ────────────────────────────────────────────────────────────────────
h2('🧩 第三章: 界面总览与核心功能')
normal(n('软件左侧为导航栏，共 8 个功能区域：'))
normal(b('🏠 主页'), n(' — 文字输入与发送，语音历史记录'))
normal(b('⚙️ 设置'), n(' — 外观、音频路由、热键、代理等基础设置'))
normal(b('🎤 语音引擎'), n(' — TTS 引擎选择与配置，角色档案管理'))
normal(b('🧩 插件'), n(' — 扩展功能入口'))
normal(b('🌍 翻译设置'), n(' — 翻译引擎与目标语言配置'))
normal(b('🎙️ 语音识别'), n(' — 语音转文字引擎与触发模式配置'))
normal(b('🎵 音频播放器'), n(' — 音乐/音效播放，支持双路输出'))
normal(b('📋 日志'), n(' — 运行日志查看，用于排查问题'))

h3('1. 🏠 主页: 文字发送')
normal(n('在文字输入框（"输入要说的话"）中输入文本。'))
normal(n('点击 '), b('发送'), n(' 按钮，或按配置的 '), b('发送热键'), n('。'))
normal(n('软件将合成语音并同时播放给你和 VRChat。'))
normal(b('自动发送: '), n('语音识别完成后自动触发发送，无需手动点击。'))
normal(b('最近识别记录: '), n('开启后，语音识别的历史记录会显示在主页下方，可点击重发。'))

h3('2. ⚙️ 设置')
normal(b('外观设置'))
normal(n('主题：跟随系统 / 浅色 / 深色'))
normal(n('强调色：自定义 UI 主色调（默认 #8A8A8A）'))
normal(n('背景图片：可设置自定义背景图'))
normal(n('背景模糊：背景虚化程度（0-100）'))
normal(b('热键设置'))
normal(b('全局唤醒热键'), n('：将窗口置顶/激活'))
normal(b('语音识别热键'), n('：触发/停止语音识别'))
normal(b('发送热键'), n('：发送当前文本框内容'))
normal(n('点击热键输入框后，按下想要的快捷键组合即可完成设置。'))
normal(b('代理（Proxy）：'), n('如需使用代理访问 API，填写 http://127.0.0.1:7890 格式的地址。'))

h3('3. 🎤 语音引擎（TTS 配置）')
normal(n('点击左侧 '), b('🎤 语音引擎'), n(' 进入 TTS 配置页。'))

h4('① Edge TTS（微软语音，推荐新手）')
normal(n('无需任何配置，开箱即用，多语言多音色支持。'))
normal(n('在 '), b('音色'), n(' 下拉框中选择声音（如 '), b('zh-CN-XiaoxiaoNeural'), n('）。'))
normal(n('可调节 '), b('语速'), n(' 和 '), b('音调'), n('。'))

h4('② GPT-SoVITS（本地克隆音色，进阶）')
normal(n('需先启动 GPT-SoVITS 后台服务（见第五章）。'))
normal(b('API 地址：'), n('默认 http://127.0.0.1:9880'))
normal(b('去除标点：'), n('✅ 使用 GPT-SoVITS 时建议开启，可解决说话"卡顿"或"长停顿"问题。'))

h4('③ FishAudio（云端 API）')
normal(n('需要填写 FishAudio 的 '), b('API Key'), n(' 和 '), b('参考音色 ID'), n('。'))

h4('角色档案管理')
normal(n('每个"角色档案"保存一套完整的 GPT-SoVITS 配置（模型路径、参考音频等）。'))
normal(b('新建'), n('：创建新档案（输入名称）'))
normal(b('重命名'), n('：修改当前档案名'))
normal(b('删除'), n('：删除当前档案'))
normal(b('切换'), n('：点击档案列表即可切换'))
normal(n('每个档案包含：GPT 模型路径（.ckpt）、SoVITS 模型路径（.pth）、参考音频路径（.wav/.mp3，3-10秒为佳）、参考文本。'))

h3('4. 🌍 翻译设置')
normal(b('翻译引擎：'))
normal(b('通用 AI 大模型（推荐）'), n('：支持智谱、DeepSeek、通义千问等兼容 OpenAI 格式的 API'))
normal(b('Google 翻译'), n('：免费，无需 Key，但网络需能访问 Google'))
normal(b('DeepL'), n('：高质量翻译，需 DeepL API Key'))
normal(b('目标语言：'), n('可配置最多 3 个目标语言，翻译结果以 | 分隔显示。'))
normal(n('例如：你好世界 | Hello World | こんにちは世界'))
normal(b('主翻译语言：'), n('设置后，TTS 将朗读翻译后的内容（例如：输入中文，说出日文）。'))

h4('📘 附：智谱 AI (BigModel) API 获取指南')
normal(n('为了使用免费且高质量的翻译，建议配置智谱 AI：'))
normal(b('第一步: 注册与登录'))
normal(n('访问智谱 AI 官网：https://www.bigmodel.cn'))
normal(n('点击右上角的 "注册/登录"，使用手机号完成注册登录 (新用户通常会赠送额度)。'))
normal(b('第二步: 获取 API Key'))
normal(n('登录成功后，把鼠标移到右上角的头像上，在下拉菜单中点击 '), b('"API密钥"'), n('。'))
normal(n('在页面右上角点击 '), b('"添加新的 API Key"'), n(' (如果没有现成的)。'))
normal(n('点击"复制"图标，复制那串长得像 2348382.sakjdhKJH... 的代码。'))
normal(b('第三步: 配置到软件中'))
normal(n('打开 '), b('[🌍 翻译设置]'), n('，引擎选择 '), b('"通用 AI 大模型"'), n('。'))
normal(n('API 地址填写：'), b('https://open.bigmodel.cn/api/paas/v4/'))
normal(n('模型名称填写：'), b('glm-4-flash'))
normal(n('粘贴 API Key，点击保存。'))
normal(b('💡 为什么推荐用智谱?'))
normal(b('免费: '), n('glm-4-flash 模型目前完全免费。'))
normal(b('懂梗: '), n('相比 Google，国产大模型对中文网络用语的理解更准确。'))
normal(b('速度快: '), n('国内服务器，无需梯子。'))

h3('5. 🎙️ 语音识别（语音转文字）')
normal(n('点击左侧 '), b('🎙️ 语音识别'), n(' 进行配置。'))
normal(b('支持的识别引擎：'))
normal(b('Windows'), n('：系统内置，无需配置，中文识别依赖系统语言包'))
normal(b('Vosk'), n('：本地离线，精度高，需下载模型（见第四章）'))
normal(b('Sherpa-ONNX'), n('：本地离线，精度最高，已内置中文模型，支持 GPU 加速'))
normal(b('触发模式：'))
normal(b('Toggle（切换）'), n('：单击热键/按钮开始识别，再次单击停止'))
normal(b('PTT（长按说话）'), n('：按住热键/按钮说话，松开即发送'))
normal(b('连续识别'), n('：软件自动检测停顿并发送，彻底解放双手'))
normal(b('Sherpa-ONNX 高级配置：'))
normal(b('运算后端'), n('：CPU（默认）/ CUDA（NVIDIA GPU）/ DML（DirectML）'))
normal(b('线程数'), n('：1-16，根据 CPU 核心数调整'))
normal(b('解码方式'), n('：greedy_search（快）/ modified_beam_search（精）'))

h3('6. 🎵 音频播放器')
normal(n('支持双路输出：你自己听音乐的同时，VRChat 里的朋友也能同步听到。'))
normal(b('支持格式：'), n('MP3、FLAC、WAV 等（内置 FFmpeg，格式通吃）'))
normal(b('用途：'), n('播放 BGM、角色语音包、整活音效等'))
normal(n('音频路由自动读取设置页面中配置的音频设备。'))

h3('7. 📱 手机输入功能')
normal(n('操作: 在主页找到二维码显示区域，用手机扫描二维码即可登录手机端控制页面。'))
normal(b('条件: '), n('需手机跟电脑连接在同一个局域网（同一个 Wi-Fi）。'))
normal(b('功能: '), n('可以在手机上输入文字发送，无需切屏。'))

blank()

# ─── 第四章 ────────────────────────────────────────────────────────────────────
h2('🔧 第四章: Vosk 离线语音模型配置（可选）')
normal(n('如果你需要使用 '), b('Vosk'), n(' 引擎（而不是内置的 Sherpa-ONNX），需要手动下载模型。'))

h3('1. 下载 Vosk 模型')
normal(n('推荐中文小模型：'), b('vosk-model-small-cn-0.22'))
normal(n('下载地址：https://alphacephei.com/vosk/models'))

h3('2. 配置步骤')
normal(n('解压下载的模型压缩包，得到一个包含 '), b('model.conf'), n(' 的文件夹。'))
normal(n('打开软件 '), b('🎙️ 语音识别'), n(' 页面，识别引擎选择 '), b('Vosk'), n('。'))
normal(n('点击 '), b('Vosk 模型路径'), n(' 右侧的 📂 图标，选中包含 model.conf 的文件夹。'))
normal(b('⚠️ 路径注意：'), n('若路径包含中文或空格，Vosk 可能加载失败。建议将模型放在全英文路径下，例如 C:\\vosk-models\\vosk-model-small-cn-0.22'))

blank()

# ─── 第五章 ────────────────────────────────────────────────────────────────────
h2('🔧 第五章: GPT-SoVITS 下载 & 后台设置 (进阶)')
normal(n('如果你想使用克隆音色（GPT-SoVITS 本地模式），必须先启动后台服务。'))

h3('1. 下载整合包')
normal(n('下载链接: https://github.com/RVC-Boss/GPT-SoVITS/releases'))
normal(n('国内直链下载1：ModelScope 镜像（搜索 GPT-SoVITS）'))
normal(n('国内直链下载2：https://hf-mirror.com（搜索 GPT-SoVITS-windows-package）'))

h3('2. 启动 API 服务 (不是 WebUI)')
normal(b('核心原则：'), n('软件连接的是 '), b('9880'), n(' 端口（API），不是网页版的 9872 端口（WebUI）。'))
normal(n('打开 GPT-SoVITS 整合包文件夹。'))
normal(n('寻找名为 '), b('go-api.bat'), n('、'), b('api.bat'), n(' 或 '), b('启动API.bat'), n(' 的文件，双击运行。'))
normal(n('如果没有这类文件，在文件夹内新建 '), b('API.bat'), n('，写入以下内容（必须放在 GPT 根目录）：'))
normal(b('@echo off'))
normal(b('runtime\\python.exe api_v2.py'))
normal(b('pause'))
normal(b('成功标志：'), n('黑色窗口显示 Uvicorn running on http://0.0.0.0:9880'))

h3('3. 准备模型文件')
normal(n('打开 '), b('🎤 语音引擎'), n(' 页面，在角色档案中填入：'))
normal(b('GPT 模型：'), n('后缀为 .ckpt 的文件'))
normal(b('SoVITS 模型：'), n('后缀为 .pth 的文件'))
normal(b('参考音频：'), n('准备一段该角色的 .wav 或 .mp3 录音（3-10秒），并记下里面的台词作为参考文本。'))

blank()

# ─── 第六章 ────────────────────────────────────────────────────────────────────
h2('❓ 第六章: 常见问题汇总 (FAQ)')

h3('Q1: 软件提示 "API连接失败" 或 "Connection Error"')
normal(b('原因 1：'), n('GPT-SoVITS 后台没开，或者开成了网页版（WebUI）。'))
normal(n('解决: 请运行 api.bat，确保黑框里显示端口是 '), b('9880'), n('。'))
normal(b('原因 2：'), n('端口号填错。'))
normal(n('解决: 检查软件设置里的 API 地址是否为 '), b('http://127.0.0.1:9880'), n('。'))
normal(b('原因 3：'), n('开了加速器/全局代理。'))
normal(n('解决: 在软件 '), b('⚙️ 设置'), n(' 的代理栏填入你的代理地址，或临时关闭系统代理。'))

h3('Q2: GPT-SoVITS 说话总是卡顿或中间长停顿')
normal(n('原因: 这是模型对标点符号过敏。'))
normal(n('解决: 在 '), b('🎤 语音引擎'), n(' 页面，开启 '), b('"去除标点"'), n(' 开关。'))

h3('Q3: GPT-SoVITS 一合成就没声音或闪退')
normal(b('原因：'), n('显卡较新，而 GPT-SoVITS 整合包内置的 CUDA 版本较老，导致兼容性冲突。'))
normal(b('解决：'), n('强制使用 CPU 进行计算。'))
normal(n('打开 GPT-SoVITS 目录 -> '), b('GPT_SoVITS/configs/tts_infer.yaml'), n('。'))
normal(n('用记事本打开，找到 '), b('device: "cuda"'), n('，改为 '), b('device: "cpu"'), n('。'))
normal(n('保存文件，重启 API 即可。'))

h3('Q4: 软件里找不到 CABLE Input 设备')
normal(n('解决: 务必 '), b('重启电脑'), n('。如果重启后还没有，请去 VBCABLE_Driver_Pack45 文件夹里重新以管理员身份运行安装程序。'))

h3('Q5: 弹窗提示 "模型加载失败: Failed to create a model"')
normal(n('原因: 语音识别引擎被设置为 Vosk，但指定路径中找不到有效的模型文件（必须包含 model.conf）。'))
normal(n('解决: 打开 '), b('🎙️ 语音识别'), n(' 页面。'))
normal(n('将识别引擎切换为 '), b('Sherpa-ONNX（内置模型）'), n(' 或 '), b('Windows'), n('。'))
normal(n('如果必须使用 Vosk，请点击 📂 重新选中直接包含 model.conf 的文件夹。'))
normal(n('检查路径是否包含中文或空格，建议改为全英文路径。'))

h3('Q6: 翻译没有生效 / 翻译显示 "NoKey"')
normal(n('解决: 在 '), b('🌍 翻译设置'), n(' 页面，确认已填入有效的 API Key。'))
normal(n('推荐使用智谱 AI（glm-4-flash 免费），配置方法见第三章。'))

h3('Q7: 热键在游戏全屏时失效')
normal(n('解决: 右键软件图标 -> 属性 -> 兼容性 -> 勾选 '), b('"以管理员身份运行此程序"'), n('，重新启动软件。'))

h3('Q8: 手机扫码后打不开网页，提示"无法连接"或"响应时间过长"')
normal(b('原因 1（最常见）：'), n('电脑网络被设置成了【公用网络】，Windows 防火墙拦截了外部设备的连接请求。'))
normal(n('解决: 点击右下角网络图标 -> 找到当前网络 -> 属性 -> 将"网络配置文件类型"从 [公用] 改为 '), b('[专用]'), n('，再重新刷新手机网页。'))
normal(b('原因 2：'), n('Windows 防火墙拦截了软件端口。'))
normal(n('解决: 搜索"允许应用通过 Windows 防火墙"，确保 '), b('Ucantalk.exe'), n(' 的【专用】和【公用】两个勾都打上。'))
normal(b('原因 3：'), n('手机和电脑没连同一个 Wi-Fi。'))
normal(n('解决: 请确保手机连接的 Wi-Fi 与电脑一致，不要使用 4G/5G 流量，也不要连接路由器的"访客网络"。'))

h3('Q9: VB-Cable 采样率报错（音频输出异常）')
normal(n('解决: 打开 Windows 系统设置 -> 系统 -> 声音 -> 更多声音设置（声音控制面板）。'))
normal(n('在"播放"选项卡中，找到 '), b('CABLE Input'), n('，右键 -> 属性 -> 高级。'))
normal(n('将"默认格式"改为 '), b('2通道，16位，44100 Hz'), n('，点击确定，重启软件。'))

h3('Q10: 语音识别 Sherpa-ONNX 引擎运行缓慢')
normal(n('若 CPU 核心数较多，可在 '), b('🎙️ 语音识别'), n(' 页面增大 '), b('线程数'), n('（如设为 4）。'))
normal(n('若有 NVIDIA 显卡，将 '), b('运算后端'), n(' 切换为 '), b('CUDA'), n('，可大幅提升速度。'))

h3('Q11: GPT-SoVITS 占用太大，VRChat 很卡')
normal(n('原因: GPT-SoVITS 默认使用 CUDA，非常吃显卡性能。'))
normal(n('解决: 同 Q3，将 tts_infer.yaml 中 device 改为 cpu（同 Q3 操作）。'))

h3('Q12: 软件没有声音，但状态显示"发送成功"')
normal(b('可能原因 1：'), n('音频路由未配置。检查 '), b('⚙️ 设置'), n(' -> '), b('音频路由'), n('，确认监听设备和 VRC 设备已正确选择。'))
normal(b('可能原因 2：'), n('VRC 麦克风未切换。确认 VRChat 内 Microphone 设置已改为 '), b('CABLE Output'), n('。'))
normal(b('可能原因 3：'), n('音量为 0。检查软件内音量设置是否为 100%。'))

blank()

# ─── 附录 ──────────────────────────────────────────────────────────────────────
h2('📁 附录: 配置文件位置')
normal(n('软件运行时的配置文件自动保存在：'))
normal(b('C:\\Users\\<你的用户名>\\AppData\\Roaming\\Ucantalk\\config.json'))
normal(n('日志文件位置：'))
normal(b('C:\\Users\\<你的用户名>\\AppData\\Roaming\\Ucantalk\\logs\\'))
normal(b('提示: '), n('如果遇到严重问题，可以尝试删除 config.json 重置所有设置为默认值。'))

blank()
normal(n('本文档适用于 Ucantalk C# 重构版（WinUI 3 / .NET 8），旧版 Python 版手册不再适用。'))

# Save
doc.save(OUTPUT)
print(f'Done! Saved to: {OUTPUT}')

